"""Tests for M1 — the parser (findr/parser.py).

User journeys:
- As the indexer, I want clean text from .txt/.md/code files so they can be
  chunked and embedded.
- As the indexer, I want text extracted from .pdf and .docx so office docs
  are searchable.
- As the indexer, I want unsupported or unreadable files skipped (None
  returned) rather than crashing the whole index run.
"""

from __future__ import annotations

import fitz  # pymupdf
import pytest
from docx import Document

from findr import parser


# --- Text / code files -------------------------------------------------------


def test_extracts_plain_text(tmp_path):
    f = tmp_path / "note.txt"
    f.write_text("the quarterly tax summary for 2025", encoding="utf-8")

    out = parser.extract_text(str(f))

    assert out is not None
    assert "quarterly tax summary" in out


def test_extracts_markdown(tmp_path):
    f = tmp_path / "readme.md"
    f.write_text("# Title\n\nSome **bold** body text.", encoding="utf-8")

    out = parser.extract_text(str(f))

    assert out is not None
    assert "Some" in out and "body text" in out


def test_extracts_source_code(tmp_path):
    f = tmp_path / "main.py"
    f.write_text("def add(a, b):\n    return a + b\n", encoding="utf-8")

    out = parser.extract_text(str(f))

    assert out is not None
    assert "def add" in out


def test_handles_non_utf8_bytes_without_crashing(tmp_path):
    f = tmp_path / "weird.txt"
    f.write_bytes(b"valid text \xff\xfe more text")

    out = parser.extract_text(str(f))

    assert out is not None
    assert "valid text" in out and "more text" in out


# --- PDF ----------------------------------------------------------------------


def test_extracts_pdf_text(tmp_path):
    f = tmp_path / "doc.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Hello from a PDF document about invoices")
    doc.save(str(f))
    doc.close()

    out = parser.extract_text(str(f))

    assert out is not None
    assert "Hello from a PDF document" in out


def test_corrupt_pdf_returns_none(tmp_path):
    f = tmp_path / "broken.pdf"
    f.write_bytes(b"%PDF-1.4 this is not a real pdf body")

    assert parser.extract_text(str(f)) is None


# --- DOCX ---------------------------------------------------------------------


def test_extracts_docx_text(tmp_path):
    f = tmp_path / "letter.docx"
    d = Document()
    d.add_paragraph("Dear team, here is the project proposal.")
    d.add_paragraph("Regards, Prince")
    d.save(str(f))

    out = parser.extract_text(str(f))

    assert out is not None
    assert "project proposal" in out
    assert "Regards, Prince" in out


def test_corrupt_docx_returns_none(tmp_path):
    f = tmp_path / "broken.docx"
    f.write_bytes(b"PK\x03\x04 not really a docx")

    assert parser.extract_text(str(f)) is None


# --- Skips / failures ---------------------------------------------------------


def test_unsupported_extension_returns_none(tmp_path):
    f = tmp_path / "photo.heic"
    f.write_bytes(b"\x00\x01\x02binary")

    assert parser.extract_text(str(f)) is None


def test_missing_file_returns_none(tmp_path):
    assert parser.extract_text(str(tmp_path / "does_not_exist.txt")) is None


def test_extension_match_is_case_insensitive(tmp_path):
    f = tmp_path / "NOTES.TXT"
    f.write_text("uppercase extension still parses", encoding="utf-8")

    out = parser.extract_text(str(f))

    assert out is not None
    assert "uppercase extension" in out
